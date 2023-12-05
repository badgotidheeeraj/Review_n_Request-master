from .models import RequestModule  # Adjust this import based on your models location
from django.http import HttpResponse, JsonResponse
from .models import UserAccount, MasterModule, RequestModule, Resource
from django.shortcuts import render, HttpResponse
from django.contrib.staticfiles import finders
from django.template.loader import get_template
from django.http import HttpResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required 
from .models import *
import base64
from django.shortcuts import get_object_or_404 
from .models import MasterModule, Resource
from django.core import serializers
from collections import OrderedDict
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.core.serializers.json import DjangoJSONEncoder
from django.core.files.storage import default_storage
from django.shortcuts import render, HttpResponse, redirect
from docx2pdf import convert
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponseRedirect
from .models import RequestModule, ReviewModule
import tempfile
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph
from docx.text.run import Run
from reportlab.lib.styles import getSampleStyleSheet
from docx.oxml import OxmlElement
import json
import os
from PyPDF2 import PdfReader, PdfWriter, PageObject
from PyPDF2.generic import AnnotationBuilder
import docx
import re
from django.core import serializers
from django.forms.models import model_to_dict
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from weasyprint import HTML
from rest_framework.response import Response
from .serializers import RequestSerializer
from rest_framework.decorators import api_view



@login_required(login_url='userLogin')
def dashboard(request):
    total_requests = RequestModule.objects.filter(username=request.user).count()
    total_reviews = ReviewModule.objects.filter(username=request.user).count()
    context = {
        'total_requests':total_requests,
        'total_reviews':total_reviews
    }
    return render(request, 'dashboard/dashboard.html',{'Title':'Dashboard', 'context':context})


''' DO NOT UPDATE COMMENTS
************************************************
*   REQUEST FUNCTIONS                          *
************************************************
'''

#  Create request
@login_required(login_url='userLogin')
def create_request(request, id):
    if request.method == 'POST':
        user_account = UserAccount.objects.get(pk=request.user.id)
        request_name = request.POST.get('request')
        ModuleName = MasterModule.objects.get(id=id)
        json_field = request.POST.get('data')

        existing_request_per_user = RequestModule.objects.filter(username=request.user, request_name=request_name).exists()

        if not existing_request_per_user:
            request_save = RequestModule(
            username=user_account, request_name=request_name, ModuleName=ModuleName, json_field=json_field) 
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request saved successfully'})
        else:
            return JsonResponse({'success': False, 'message': 'Request name already exist'})

    selectedModule = MasterModule.objects.get(id=id)
    json_data = json.loads(selectedModule.json_field) 

    generalPoint = json_data['General Point']
    technicalPoint = json_data['Technical Point']

    resources = Resource.objects.get(id=1).json_field

    return render(request, "dashboard/create_request.html", {'Title':'Request','selectedModule': selectedModule,'standardmaster':'master ', 'generalPoint': generalPoint, 'technicalPoint': technicalPoint, 'resources': resources})

# Show all request
@login_required(login_url='userLogin')
@api_view(['GET'])
def show_all_request(request):
    modulesList = MasterModule.objects.all()
    page = request.GET.get('page')
    fav = request.GET.get('fav')
    
    if fav == 'true':
        all_request = RequestModule.objects.filter(
            username=request.user, is_favorite=True).order_by('-created_at')
    else:
        all_request = RequestModule.objects.filter(
            username=request.user).order_by('-created_at')
    items_per_page = 3
    paginator = Paginator(all_request, items_per_page) 
    # ReviewModule
    # reviewList
    page = request.GET.get('page')

    if request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest':
        try:
            all_request = paginator.page(page)
        except PageNotAnInteger:
            all_request = paginator.page(1)
        except EmptyPage:
            all_request = paginator.page(paginator.num_pages)

        serialized_request_list = RequestSerializer(all_request, many=True)
        
        all_pages = all_request.paginator.num_pages
        current_page = all_request.number

        has_previous = all_request.has_previous()
        if has_previous:
            previous_page_number = all_request.previous_page_number()
        else:
            previous_page_number = False

        has_next = all_request.has_next()
        if has_next:
            next_page_number = all_request.next_page_number()
        else:
            next_page_number = False
            
        data = {
            'requestList': serialized_request_list.data,
            'has_previous': has_previous,
            'has_next': has_next,
            'previous_page_number': previous_page_number,
            'next_page_number': next_page_number,     
            'all_pages': all_pages,     
            'current_page': current_page,     
        }    
  
        return Response(data)
        # return JsonResponse(data)

    return render(request, 'dashboard/show_all_request.html', {'Title':'Request',"master": modulesList})


# Update request
@login_required(login_url='userLogin')
def update_request(request, id):
    if request.method == 'POST':
        request_name = request.POST.get('request')   
        originalName = request.POST.get('originalName')
        json_field1 = request.POST.get('data')
        existing_request = RequestModule.objects.filter(username=request.user, request_name=request_name).exists()

        if request_name == originalName:
            request_save = RequestModule.objects.get(id=id)
            request_save.json_field = json_field1
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request Updated successfully'})
        else:
            existing_request = RequestModule.objects.filter(username=request.user, request_name=request_name).exists()


        if not existing_request:
            request_save = RequestModule.objects.get(id=id)
            request_save.request_name = request_name
            request_save.json_field = json_field1
            request_save.save()
            return JsonResponse({'success': True, 'message': 'Request Update successfully'})
        else:
            return JsonResponse({'success': False, 'message': 'Request name already exist'})   

    master = RequestModule.objects.get(id=id)
    json_data = json.loads(master.json_field)

    generalPoint = json_data['General Point']
    technicalPoint = json_data['Technical Point']

    resources = Resource.objects.get(id=1).json_field
        # return JsonResponse(True)
    return render(request, 'dashboard/update_request.html', {'Title':'Request','standard': master, 'generalPoint': generalPoint, 'technicalPoint': technicalPoint, 'resources': resources})

# Delete request
@login_required(login_url='userLogin')
def delete_request(request, id):
    RequestModule.objects.get(id=id).delete()
    return redirect('show_all_request')

# Delete all requests

@csrf_exempt
def delete_all_request(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_ids[]')
        print(selected_ids)
        RequestModule.objects.filter(
        username=request.user, id__in=selected_ids).delete()
    return JsonResponse({'success': True, 'message': 'Request Delete successfully'})

# Generate request Pdf  
@csrf_exempt
def create_request_pdf(request):
    if request.method=='POST':
        id=request.POST.get('id')
        temp = RequestModule.objects.get(id=id) 
        data = json.loads(temp.json_field)
        requestname = temp.request_name
        context = {
        "Technical": data["Technical Point"], 
        "General": data["General Point"],
        'requestname': requestname,
        "clientname": request.user.username,
        "modulename": temp.ModuleName,
        "date": temp.created_at
        }
        html_string = render_to_string('dashboard/pdf.html', context)

        pdf = HTML(string=html_string).write_pdf()

        pdf_base64 = base64.b64encode(pdf).decode('utf-8')

        # response = HttpResponse(pdf_base64, content_type='application/pdf')
        response = HttpResponse(pdf_base64, content_type='text/plain')
        # response['Content-Disposition'] = 'attachment; filename="sample.pdf"'
        return response






''' DO NOT UPDATE COMMENTS
************************************************
*   REVIEW FUNCTIONS                          *
************************************************
'''

# 1. Create review - Master (formally admin/default created) 
@login_required(login_url='userLogin')
def create_master_review(request):
    if request.method == 'POST':
        docx_file = request.FILES['docx_file']
        request_name = request.POST.get('request_name',)


        checkpoints = MasterModule.objects.get(id=request_name).json_field
        json_data = json.loads(checkpoints)
        general = json_data['General Point']
        technical = json_data['Technical Point']

        exclude_keys = ["General Point", "Technical Point"] 
        keys = []
        def extract_keys_recursive(dictionary):
            for key, value in dictionary.items():
                keys.append(key)  # Include the key itself
                if isinstance(value, dict):
                    extract_keys_recursive(value)  # Recurse into nested dictionary 
        extract_keys_recursive(json_data)
        exclude_keys = ["General Point", "Font", "Page Layout"]
        filtered_keys = [key for key in keys if key not in exclude_keys]
        word_file = docx_file  
        # Define the keys and filtered_keys based on your requirements
        filtered_keys = ['Font Style', 'Font Size', 'Font Color', 'Background Color', 'Bold', 'Italic', 'Underline', 'Page Column', 'Spacing Before', 'Spacing After', 'Margin Top', 'Margin Left', 'Margin Right', 'Margin Bottom', 'Indent Right', 'Indent Left', 'Alignment', 'Words like we, I, our and you']
        def analyze_word_document(word_file, filtered_keys):
            # Load the Word document
            doc = docx.Document(word_file)
            # Initialize a dictionary to store property values
            properties = {}
            # Define the words to check for
            words_to_check = ["we", "i", "our", "you"]
            # Check if any of the words to check for are present in the document
            found_words = set()
            for paragraph in doc.paragraphs:
                for word in words_to_check:
                    if word in paragraph.text.lower():
                        found_words.add(word)
            # Iterate through the JSON data to fetch property names and values
            category_dict = OrderedDict()
            for property_name in filtered_keys:
                # Initialize the set for this property
                property_set = set()
                # Handle "Indent Right" and "Indent Left" properties
                if property_name == "Indent Right" or property_name == "Indent Left":
                    value = None
                    for paragraph in doc.paragraphs:
                        if property_name == "Indent Right":
                            value = paragraph.paragraph_format.right_indent
                        elif property_name == "Indent Left":
                            value = paragraph.paragraph_format.left_indent
                        # Handle the case where the value is not None
                        if value is not None:
                            property_set.add(str(value.pt))  # Convert to string
                # Handle other properties
                else:
                    # Iterate through paragraphs and runs to collect unique formatting values
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            # Handle the specific property
                            if property_name == "Font Style":
                                property_set.add(font.name)
                            elif property_name == "Font Size" and font.size and font.size.pt:
                                property_set.add(str(int(font.size.pt)))  # Convert to string
                            elif property_name == "Bold":
                                property_set.add("Yes" if font.bold else "No")
                            elif property_name == "Italic":
                                property_set.add("Yes" if font.italic else "No")
                            elif property_name == "Underline":
                                property_set.add("Yes" if font.underline else "No")
                            elif property_name == "Font Color" and font.color:
                                # Check if font color has RGB and get the RGB string
                                if hasattr(font.color.rgb, 'rgb_str'):
                                    property_set.add(font.color.rgb.rgb_str())
                                else:
                                    property_set.add("None")
                            elif property_name == "Background Color":
                                # Check if run has shading and if shading has RGB
                                if hasattr(run, 'shading') and hasattr(run.shading, 'background_color') and hasattr(run.shading.background_color, 'rgb_str'):
                                    property_set.add(run.shading.background_color.rgb_str())
                                else:
                                    property_set.add("None")
                            # Add more conditions for other properties as needed
                    # Special handling for additional properties
                    if property_name == "Page Column":
                        property_set.add("1")
                    if property_name in ["Spacing Before", "Spacing After", "Margin Top", "Margin Left",
                                        "Margin Right", "Margin Bottom"]:
                        for paragraph in doc.paragraphs:
                            value = None
                            if property_name == "Spacing Before":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Spacing After":
                                value = paragraph.paragraph_format.space_after
                            elif property_name == "Margin Top":
                                value = paragraph.paragraph_format.space_before
                            elif property_name == "Margin Left":
                                value = paragraph.paragraph_format.left_indent
                            elif property_name == "Margin Right":
                                value = paragraph.paragraph_format.right_indent
                            elif property_name == "Margin Bottom":
                                value = paragraph.paragraph_format.space_after

                            # Handle the case where the value is not None
                            if value is not None:
                                property_set.add(str(value.pt))  # Convert to string
                    if property_name == "Alignment":
                        for paragraph in doc.paragraphs:
                            property_set.add(str(paragraph.alignment))  # Convert to string
                    if property_name == "Words like we, I, our and you":
                        if found_words:
                            property_set.add("Yes")
                        else:
                            property_set.add("No")
                # Convert sets to lists for JSON serialization
                properties[property_name] = list(property_set)
            return properties
            
        # Define the format_json function
        def format_json(data, level=0):
            result = ""
            if isinstance(data, dict):
                result += "{\n"
                for key, value in data.items():
                    result += '    ' * level + f'"{key}": '
                    result += format_json(value, level + 1)
                    result += ",\n"
                result = result[:-2]  # Remove the trailing comma and newline
                result += "\n" + '    ' * (level - 1) + "}"
            elif isinstance(data, list):
                result += "["
                for item in data:
                    result += format_json(item, level + 1) + ", "
                result = result[:-2]  # Remove the trailing comma and space
                result += "]"      
            else:
                result += f'"{data}"'
            return result

        # Call the analyze_word_document function
        result = analyze_word_document(word_file, filtered_keys)

        # Format the JSON result
        formatted_result = json.loads(format_json(result))


        # fake comment pool
        comment_pool = {
            "Font Style": "comment of Font Style",
            "Font Size": " comment of Font Size",
            "Font Color": 'comment of Font Color',
            "Background Color": "comment of Background Color",
            "Bold": 'comment of Bold',
            "Italic": 'comment of  Italic',
            "Underline": 'comment of Underline',
            "Page Column": 'comment of Page Column',
            "Margin Top": 'comment of Margin Top',
            "Margin Left": 'comment of Margin Left',
            "Margin Right": 'comment of Margin Right',
            "Margin Bottom": 'comment of Margin Bottom',
            "Indent Right": 'comment of Indent Right',
            "Indent Left": 'comment of Indent Left',
            "Spacing Before": 'comment of Spacing Before',
            "Spacing After": 'comment of Spacing After',
            "Alignment": 'comment of Alignment',
            "Total Words":'comment of Total Word',
            "Document Type":'comment of Document Type',
            "Accepted Grammar Mistake":'comment of Accepted Grammar Mistake',
            "Accepted Plagiarism":'comment of Accepted Plagiarism',
            "Plag Tool":'comment of Plag Tool',
            "Words like we, I, our and you":'comment of Words like we, I, our and you',
        }

        # fake img pool
        review_img_pool = {
            "Font": ["login-main1.png"],
        }

        return render(request, 'dashboard/create_review.html', {'Title':'Review','review_type': 'Standard','document_path': docx_file, 'General': general, 'Technical': technical, "fake_data": formatted_result, 'requstDetail': request_name, 'fake_img':review_img_pool,'review_comment':comment_pool,})


# 2. Create review - Custom (formally user created, derived from Master review)
@login_required(login_url='userLogin')
def create_custom_review(request):
    if request.method == 'POST':
        docx_file = request.FILES['docx_file'] 
        request_name = request.POST.get('request_name',) 
        req = RequestModule.objects.get(id=request_name) 
        checkpoints = req.json_field  
        json_data = json.loads(checkpoints)
        general = json_data['General Point']
        technical = json_data['Technical Point']

        comment_pool = {
            "Font Style": "comment of Font Style",
            "Font Size": " comment of Font Size",
            "Font Color": 'comment of Font Color',
            "Background Color": "comment of Background Color",
            "Bold": 'comment of Bold',
            "Italic": 'comment of  Italic',
            "Underline": 'comment of Underline',
            "Page Column": 'comment of Page Column',
            "Margin Top": 'comment of Margin Top',
            "Margin Left": 'comment of Margin Left',
            "Margin Right": 'comment of Margin Right',
            "Margin Bottom": 'comment of Margin Bottom',
            "Indent Right": 'comment of Indent Right',
            "Indent Left": 'comment of Indent Left',
            "Spacing Before": 'comment of Spacing Before',
            "Spacing After": 'comment of Spacing After',
            "Alignment": 'comment of Alignment',
            "Total Words":'comment of Total Word',
            "Document Type":'comment of Document Type',
            "Accepted Grammar Mistake":'comment of Accepted Grammar Mistake',
            "Accepted Plagiarism":'comment of Accepted Plagiarism',
            "Plag Tool":'comment of Plag Tool',
            "Words like we, I, our and you":'comment of Words like we, I, our and you',
        }


        # det_font = []
        det_font1 = {
            "Font Style": ['Times New Roman'],
            "Font Size": ['20', '30'],
            "Font Color": ['Black', 'Yellow'],
            "Background Color": ['Black'],
            "Bold": ['yes'],
            "Italic": ['No'],
            "Underline": ['No'],
            "Page Column": ['One Column'],
            "Margin-Top": ['1.2'],
            "Margin-Left": ['1.0'],
            "Margin-Right": ['1.4'],
            "Margin-Bottom": ['1.0'],
            "Indent-Right": ['1.4'],
            "Indent-Left": ['1.8'],
            "Spacing-Before": ['1.0'],
            "Spacing-After": ['1'],
            "Alignment": ['Center'],
            "Total Words":["150-200"] 
        }

    return render(request, 'dashboard/create_review.html', {'Title':'Review','review_type': 'Custom','document_path': docx_file, 'General': general, 'Technical': technical, "fake_data": det_font1, 'requstDetail': request_name,"req_name":req,'review_comment':comment_pool})


'''
Get the all custom requests created 
by user associated with the selected module.
'''
@login_required(login_url='userLogin')
def getCustomRequestForReview(request, id):
    request_with_module = RequestModule.objects.filter(username=request.user, ModuleName_id = id)
    data = serializers.serialize('json', request_with_module)
    return JsonResponse({'success': True, 'message': 'Request saved successfully', 'data': data})   


# Save review explicitly
@csrf_exempt
def save_review_report(request):
    if request.method == 'POST':
        try:
            request_name = request.POST.get('request_name')
            docx_file = request.POST.get('docx_file')  # Use request.FILES.get() to handle file upload  
            processed_file = request.POST.get('processed_file')
            json_field = request.POST.get('json_field')  
            review_type = request.POST.get('reviewtype')  
            review_name = request.POST.get('review_name')
            existing_review = ReviewModule.objects.filter(username=request.user, review_name=review_name).exists()

            if not existing_review:
                document = ReviewModule(username=request.user,request_name=request_name,docx_file=docx_file,processed_file=processed_file,json_field=json.loads( json_field), review_type =review_type, review_name =review_name)
                document.save()
                return JsonResponse({"success": True, "message": "Review saved successfully"}, status=200)
            else:
                return JsonResponse({'success': False, 'message': 'Review name already exist'})

        except Exception as e:
            return JsonResponse({"error": f"Error processing data: {str(e)}"}, status=500)
    else:
        return JsonResponse({"error": "Invalid request method"}, status=405)
 


# Show All Reviews
@login_required(login_url='userLogin')
def show_all_review(request):
    modulesList = MasterModule.objects.all()
    page = request.GET.get('page')
    fav = request.GET.get('fav')
    if fav == 'true':
        reviewList = ReviewModule.objects.filter(
            username=request.user, is_favorite=True).order_by('-created_at')
    else:
        reviewList = ReviewModule.objects.filter(
            username=request.user).order_by('-created_at')
    items_per_page = 3
    paginator = Paginator(reviewList, items_per_page)
    page = request.GET.get('page')

    if request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest':
        try:
            reviewList = paginator.page(page)
        except PageNotAnInteger:
            reviewList = paginator.page(1)
        except EmptyPage:
            reviewList = paginator.page(paginator.num_pages)

        serialized_review_list = serializers.serialize('json', reviewList)
        all_pages = reviewList.paginator.num_pages
        current_page = reviewList.number

        has_previous = reviewList.has_previous()
        if has_previous:
            previous_page_number = reviewList.previous_page_number()
        else:
            previous_page_number = False

        has_next = reviewList.has_next()
        if has_next:
            next_page_number = reviewList.next_page_number()
        else:
            next_page_number = False


        data = {
        'reviewList': serialized_review_list, 
        'has_previous': has_previous,
        'has_next': has_next,
        'previous_page_number': previous_page_number,
        'next_page_number': next_page_number,        
        'all_pages': all_pages,
        'current_page': current_page,
        
        }    

        return JsonResponse(data)
    return render(request, 'dashboard/show_all_review.html', {'Title':'Review', 'modulesList': modulesList ,"reviewList":reviewList })


# Delete review
@login_required(login_url='userLogin')
def delete_review(request, id):
    ReviewModule.objects.get(id=id).delete()
    return redirect('show_all_review')

# Delete all review
@csrf_exempt
@login_required
def delete_all_review(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_ids[]')
        ReviewModule.objects.filter(username=request.user, id__in=selected_ids).delete()

        return JsonResponse({'message': 'Reviews deleted successfully'})
    return JsonResponse({'error': 'Invalid request method'}, status=400)


# Update review
@login_required(login_url='userLogin')
def update_review(request, review_id):
    if request.method =='POST':
        originalName = request.POST.get('originalName')
        json_field = request.POST.get('json_field')
        review_name = request.POST.get('review_name')

        obj = ReviewModule.objects.get(id = review_id) 
        if review_name == originalName:
            obj.json_field = json.loads(json_field)
            obj.save()
            return JsonResponse({'success': True, 'message': 'Review Updated successfully'})
        else:
            existing_review = ReviewModule.objects.filter(username=request.user, review_name=review_name).exists()
            if existing_review == False:
                obj.review_name = review_name
                obj.json_field = json.loads(json_field)
                obj.save()
                return JsonResponse({'success': True, "message": "Data Updated successfully"}, status=200)
            else:
                return JsonResponse({'success': False, 'message': 'Review name already exist'})
    
    review = ReviewModule.objects.get(id = review_id)
    review_img_pool = {
        "Font": ["login-main1.png"],
        }

    return render(request, 'dashboard/update_review.html', {'Title': 'Review', 'fake_img': review_img_pool, 'review': review})

 
@csrf_exempt
def create_review_pdf(request):
    if request.method=="POST":
        id=request.POST.get('id')
        json_data = ReviewModule.objects.get(id=id).json_field
        req_name = ReviewModule.objects.get(id=id)
        rev_type = ReviewModule.objects.get(id=id).review_type
        context = {
        'data': json_data,
        'req_name': req_name,
        'rev_type': rev_type,}
        html_string = render_to_string('dashboard/reviewpdf.html', context)
        pdf = HTML(string=html_string).write_pdf()
        pdf_base64 = base64.b64encode(pdf).decode('utf-8')
        response = HttpResponse(pdf_base64, content_type='application/pdf')
        return response


@csrf_exempt
def add_favorite_request(request):
    try:
        if request.method =="POST":
            id = request.POST.get('id')
            condition = request.POST.get('isFavorite')
            favorite = condition.lower() == 'false'
            if favorite != True:
                fav_instance = RequestModule.objects.get(id=id)
                fav_instance.is_favorite = True
                fav_instance.save()
                return JsonResponse({"message": "add to Favorite",'status':201})
            else:
                fav_instance = RequestModule.objects.get(id=id)
                fav_instance.is_favorite = False
                fav_instance.save()
                return JsonResponse({"message": "remove from Favorite",'status':200})

    except Exception as e:
        return JsonResponse({'error':'An unexpected error occurred'},status=500)


@csrf_exempt
def add_favorite_review(request):
    try:
        if request.method == "POST":
            id = request.POST.get('id')
            condition = request.POST.get('isFavorite')
            favorite = condition.lower() == 'false'
            if favorite != True:
                fav_instance = ReviewModule.objects.get(id=id)
                fav_instance.is_favorite = True
                fav_instance.save()
                return JsonResponse({"message": "add to Favorite", 'status': 201})
            else:
                fav_instance = ReviewModule.objects.get(id=id)
                fav_instance.is_favorite = False
                fav_instance.save()
                return JsonResponse({"message": "remove from Favorite", 'status': 200})

    except Exception as e:
        return JsonResponse({'error': 'An unexpected error occurred'}, status=500)